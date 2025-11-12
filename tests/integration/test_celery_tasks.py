"""
Integration tests for Celery tasks.

Tests the complete Celery task execution including:
- Task queueing
- Redis communication
- Task execution
- Result retrieval
- Error handling and retries
"""

import pytest
from unittest.mock import patch, MagicMock
from celery.exceptions import Retry

# Import tasks to test
from tasks import (
    process_pdf_task,
    process_excel_task,
    process_word_task,
    transcribe_audio_task,
    scrape_url_task,
)


@pytest.mark.integration
@pytest.mark.database
class TestPDFProcessingTask:
    """Test PDF processing Celery task."""

    @pytest.mark.asyncio
    async def test_process_pdf_success(self, db_session, tmp_path):
        """
        Test successful PDF processing.
        """
        # Create a simple test PDF file
        pdf_path = tmp_path / "test.pdf"
        pdf_path.write_text("Test PDF content")  # Simplified for test

        # Create test user and document in DB
        from database.models import User, Document

        user = User(telegram_id=12345, username="testuser")
        db_session.add(user)
        db_session.flush()

        document = Document(
            user_id=user.id,
            filename="test.pdf",
            file_path=str(pdf_path),
            status="pending"
        )
        db_session.add(document)
        db_session.commit()

        # Execute task (mock actual PDF processing)
        with patch('tasks.fitz.open') as mock_fitz:
            mock_doc = MagicMock()
            mock_doc.load_page.return_value.get_text.return_value = "PDF text content"
            mock_doc.page_count = 1
            mock_fitz.return_value = mock_doc

            with patch('tasks.get_db_session', return_value=db_session):
                result = process_pdf_task(document.id)

        # Verify task result
        assert result is not None

        # Verify document status updated
        db_session.refresh(document)
        assert document.status == "completed"
        assert document.text_content is not None

    @pytest.mark.asyncio
    async def test_process_pdf_retry_on_error(self, db_session, tmp_path):
        """
        Test PDF processing retries on error.
        """
        pdf_path = tmp_path / "test.pdf"
        pdf_path.write_text("Test")

        from database.models import User, Document

        user = User(telegram_id=12345, username="testuser")
        db_session.add(user)
        db_session.flush()

        document = Document(
            user_id=user.id,
            filename="test.pdf",
            file_path=str(pdf_path),
            status="pending"
        )
        db_session.add(document)
        db_session.commit()

        # Simulate error that triggers retry
        with patch('tasks.fitz.open', side_effect=IOError("File error")):
            with patch('tasks.get_db_session', return_value=db_session):
                with pytest.raises((Retry, IOError)):
                    process_pdf_task(document.id)


@pytest.mark.integration
@pytest.mark.database
class TestExcelProcessingTask:
    """Test Excel processing Celery task."""

    @pytest.mark.asyncio
    async def test_process_excel_success(self, db_session, tmp_path):
        """
        Test successful Excel processing.
        """
        # Create test Excel file
        import pandas as pd

        excel_path = tmp_path / "test.xlsx"
        df = pd.DataFrame({
            'Name': ['Alice', 'Bob', 'Charlie'],
            'Value': [100, 200, 300]
        })
        df.to_excel(excel_path, index=False)

        from database.models import User, Document

        user = User(telegram_id=12345, username="testuser")
        db_session.add(user)
        db_session.flush()

        document = Document(
            user_id=user.id,
            filename="test.xlsx",
            file_path=str(excel_path),
            status="pending"
        )
        db_session.add(document)
        db_session.commit()

        # Execute task
        with patch('tasks.get_db_session', return_value=db_session):
            result = process_excel_task(document.id)

        # Verify result
        assert result is not None

        # Verify document processed
        db_session.refresh(document)
        assert document.status == "completed"
        assert "Alice" in document.text_content or "100" in document.text_content


@pytest.mark.integration
@pytest.mark.database
class TestWordProcessingTask:
    """Test Word document processing Celery task."""

    @pytest.mark.asyncio
    async def test_process_word_success(self, db_session, tmp_path):
        """
        Test successful Word document processing.
        """
        from docx import Document as DocxDocument

        # Create test Word document
        word_path = tmp_path / "test.docx"
        doc = DocxDocument()
        doc.add_paragraph("Test paragraph 1")
        doc.add_paragraph("Test paragraph 2")
        doc.save(str(word_path))

        from database.models import User, Document

        user = User(telegram_id=12345, username="testuser")
        db_session.add(user)
        db_session.flush()

        document = Document(
            user_id=user.id,
            filename="test.docx",
            file_path=str(word_path),
            status="pending"
        )
        db_session.add(document)
        db_session.commit()

        # Execute task
        with patch('tasks.get_db_session', return_value=db_session):
            result = process_word_task(document.id)

        # Verify result
        assert result is not None

        # Verify document processed
        db_session.refresh(document)
        assert document.status == "completed"
        assert "Test paragraph" in document.text_content


@pytest.mark.integration
@pytest.mark.slow
class TestAudioTranscriptionTask:
    """Test audio transcription Celery task."""

    @pytest.mark.asyncio
    async def test_transcribe_audio_mock(self, db_session, tmp_path):
        """
        Test audio transcription with mocked OpenAI API.
        """
        # Create dummy audio file
        audio_path = tmp_path / "test.mp3"
        audio_path.write_bytes(b"fake audio data")

        from database.models import User, Document

        user = User(telegram_id=12345, username="testuser")
        db_session.add(user)
        db_session.flush()

        document = Document(
            user_id=user.id,
            filename="test.mp3",
            file_path=str(audio_path),
            status="pending"
        )
        db_session.add(document)
        db_session.commit()

        # Mock OpenAI transcription
        with patch('tasks.openai.Audio.transcribe') as mock_transcribe:
            mock_transcribe.return_value = MagicMock(
                text="This is the transcribed audio content."
            )

            with patch('tasks.get_db_session', return_value=db_session):
                result = transcribe_audio_task(document.id)

        # Verify result
        assert result is not None

        # Verify document processed
        db_session.refresh(document)
        assert document.status == "completed"
        assert "transcribed" in document.text_content.lower()


# TODO: Implement query_document_task and uncomment these tests
# @pytest.mark.integration
# @pytest.mark.database
# class TestAIQueryTask:
#     """Test AI query Celery task."""
#
#     @pytest.mark.asyncio
#     async def test_query_document_success(self, db_session):
#         """
#         Test successful AI query execution.
#         """
#         from database.models import User, Document
#
#         user = User(telegram_id=12345, username="testuser")
#         db_session.add(user)
#         db_session.flush()
#
#         document = Document(
#             user_id=user.id,
#             filename="test.pdf",
#             text_content="This document contains sales data for Q1 2024.",
#             status="completed"
#         )
#         db_session.add(document)
#         db_session.commit()
#
#         # Execute AI query task with mocked Gemini
#         with patch('tasks.genai.GenerativeModel') as mock_model:
#             mock_response = MagicMock()
#             mock_response.text = "The document shows Q1 2024 sales data."
#             mock_model.return_value.generate_content.return_value = mock_response
#
#             with patch('tasks.get_db_session', return_value=db_session):
#                 result = query_document_task(
#                     document_id=document.id,
#                     query="What does this document contain?",
#                     user_id=user.id
#                 )
#
#         # Verify AI response
#         assert result is not None
#         assert "Q1 2024" in result or "sales" in result.lower()
#
#     @pytest.mark.asyncio
#     async def test_query_with_caching(self, db_session, mock_redis_client):
#         """
#         Test AI query with caching enabled.
#         """
#         from database.models import User, Document
#
#         user = User(telegram_id=12345, username="testuser")
#         db_session.add(user)
#         db_session.flush()
#
#         document = Document(
#             user_id=user.id,
#             filename="test.pdf",
#             text_content="Sales report Q1",
#             status="completed"
#         )
#         db_session.add(document)
#         db_session.commit()
#
#         query = "What is in the document?"
#
#         # First query - should hit AI
#         with patch('tasks.genai.GenerativeModel') as mock_model:
#             mock_response = MagicMock()
#             mock_response.text = "This is a sales report for Q1."
#             mock_model.return_value.generate_content.return_value = mock_response
#
#             with patch('tasks.get_db_session', return_value=db_session):
#                 with patch('tasks.redis_client', mock_redis_client):
#                     result1 = query_document_task(document.id, query, user.id)
#
#             first_call_count = mock_model.return_value.generate_content.call_count
#
#         # Second identical query - should use cache
#         with patch('tasks.genai.GenerativeModel') as mock_model:
#             with patch('tasks.get_db_session', return_value=db_session):
#                 with patch('tasks.redis_client', mock_redis_client):
#                     # Set cache hit
#                     mock_redis_client.get.return_value = result1
#                     result2 = query_document_task(document.id, query, user.id)
#
#             second_call_count = mock_model.return_value.generate_content.call_count
#
#         # If caching works, AI should not be called second time
#         # assert second_call_count == 0  # Depends on implementation


@pytest.mark.integration
@pytest.mark.redis
class TestTaskWithRedis:
    """Test tasks that depend on Redis."""

    def test_task_queuing(self):
        """
        Test that tasks are properly queued in Redis.
        """
        pytest.skip("Requires real Celery worker and Redis")

    def test_task_result_retrieval(self):
        """
        Test retrieving task results from Redis.
        """
        pytest.skip("Requires real Celery worker and Redis")


@pytest.mark.integration
@pytest.mark.slow
class TestTaskPerformance:
    """Test task performance and timeouts."""

    def test_task_timeout(self):
        """
        Test that long-running tasks timeout correctly.
        """
        pytest.skip("Performance test - run manually")

    def test_concurrent_tasks(self):
        """
        Test multiple tasks running concurrently.
        """
        pytest.skip("Performance test - run manually")


if __name__ == "__main__":
    """
    Run integration tests for Celery tasks.

    Usage:
        pytest tests/integration/test_celery_tasks.py -v
        pytest tests/integration/test_celery_tasks.py -v -m "integration and not slow"
    """
    pytest.main([__file__, "-v", "--tb=short"])
