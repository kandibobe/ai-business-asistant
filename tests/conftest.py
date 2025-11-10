"""
Pytest configuration and fixtures.
Shared fixtures for all tests.
"""
import os
import pytest
import tempfile
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker, Session
from sqlalchemy.pool import StaticPool
from faker import Faker

# Установить тестовый режим
os.environ['TESTING'] = 'true'

from database.models import Base, User, Document
from database.database import get_db


# ============================================================================
# Database Fixtures
# ============================================================================

@pytest.fixture(scope="function")
def db_engine():
    """
    Create in-memory SQLite database engine for testing.
    Each test gets a fresh database.
    """
    engine = create_engine(
        "sqlite:///:memory:",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    Base.metadata.create_all(bind=engine)
    yield engine
    Base.metadata.drop_all(bind=engine)
    engine.dispose()


@pytest.fixture(scope="function")
def db_session(db_engine):
    """
    Create a new database session for a test.
    Automatically rolls back after test completes.
    """
    TestingSessionLocal = sessionmaker(
        autocommit=False,
        autoflush=False,
        bind=db_engine
    )
    session = TestingSessionLocal()
    yield session
    session.rollback()
    session.close()


# ============================================================================
# Data Fixtures
# ============================================================================

@pytest.fixture
def faker_instance():
    """Faker instance for generating test data."""
    return Faker(['ru_RU', 'en_US'])


@pytest.fixture
def sample_user_data(faker_instance):
    """Sample user data for testing."""
    return {
        'user_id': faker_instance.random_int(min=10000, max=99999999),
        'username': faker_instance.user_name(),
        'first_name': faker_instance.first_name(),
        'last_name': faker_instance.last_name(),
    }


@pytest.fixture
def sample_user(db_session, sample_user_data):
    """Create a sample user in the database."""
    user = User(**sample_user_data)
    db_session.add(user)
    db_session.commit()
    db_session.refresh(user)
    return user


@pytest.fixture
def sample_document_data(faker_instance):
    """Sample document data for testing."""
    return {
        'file_name': faker_instance.file_name(extension='pdf'),
        'file_path': faker_instance.file_path(depth=3, extension='pdf'),
        'content': faker_instance.text(max_nb_chars=1000),
        'document_type': 'pdf',
        'file_size': faker_instance.random_int(min=1000, max=1000000),
        'word_count': faker_instance.random_int(min=100, max=10000),
        'char_count': faker_instance.random_int(min=500, max=50000),
    }


@pytest.fixture
def sample_document(db_session, sample_user, sample_document_data):
    """Create a sample document in the database."""
    document = Document(
        user_id=sample_user.id,
        **sample_document_data
    )
    db_session.add(document)
    db_session.commit()
    db_session.refresh(document)
    return document


# ============================================================================
# File Fixtures
# ============================================================================

@pytest.fixture
def temp_directory():
    """Create a temporary directory for test files."""
    with tempfile.TemporaryDirectory() as tmpdir:
        yield tmpdir


@pytest.fixture
def sample_pdf_file(temp_directory, faker_instance):
    """Create a sample PDF file for testing."""
    import fitz  # PyMuPDF

    pdf_path = os.path.join(temp_directory, "test_document.pdf")

    # Create a simple PDF with text
    doc = fitz.open()
    page = doc.new_page()
    text = faker_instance.text(max_nb_chars=500)
    page.insert_text((72, 72), text)
    doc.save(pdf_path)
    doc.close()

    return pdf_path


@pytest.fixture
def sample_excel_file(temp_directory, faker_instance):
    """Create a sample Excel file for testing."""
    import pandas as pd

    excel_path = os.path.join(temp_directory, "test_spreadsheet.xlsx")

    # Create sample data
    data = {
        'Name': [faker_instance.name() for _ in range(10)],
        'Age': [faker_instance.random_int(min=18, max=80) for _ in range(10)],
        'City': [faker_instance.city() for _ in range(10)],
        'Salary': [faker_instance.random_int(min=30000, max=150000) for _ in range(10)],
    }
    df = pd.DataFrame(data)
    df.to_excel(excel_path, index=False)

    return excel_path


@pytest.fixture
def sample_word_file(temp_directory, faker_instance):
    """Create a sample Word file for testing."""
    from docx import Document

    word_path = os.path.join(temp_directory, "test_document.docx")

    doc = Document()
    doc.add_heading(faker_instance.sentence(), 0)

    for _ in range(5):
        doc.add_paragraph(faker_instance.paragraph())

    # Add a table
    table = doc.add_table(rows=3, cols=3)
    for row in table.rows:
        for cell in row.cells:
            cell.text = faker_instance.word()

    doc.save(word_path)

    return word_path


@pytest.fixture
def sample_audio_file(temp_directory):
    """Create a sample audio file for testing."""
    from pydub import AudioSegment
    from pydub.generators import Sine

    audio_path = os.path.join(temp_directory, "test_audio.mp3")

    # Generate 1 second of 440 Hz sine wave
    sine_wave = Sine(440).to_audio_segment(duration=1000)
    sine_wave.export(audio_path, format="mp3")

    return audio_path


# ============================================================================
# Mock Fixtures
# ============================================================================

@pytest.fixture
def mock_telegram_update(faker_instance):
    """Mock Telegram Update object."""
    from unittest.mock import MagicMock

    update = MagicMock()
    update.effective_user.id = faker_instance.random_int(min=10000, max=99999999)
    update.effective_user.username = faker_instance.user_name()
    update.effective_user.first_name = faker_instance.first_name()
    update.effective_user.last_name = faker_instance.last_name()
    update.message.text = faker_instance.sentence()
    update.message.chat_id = faker_instance.random_int(min=10000, max=99999999)

    return update


@pytest.fixture
def mock_telegram_context():
    """Mock Telegram Context object."""
    from unittest.mock import MagicMock

    context = MagicMock()
    return context


@pytest.fixture
def mock_gemini_model():
    """Mock Google Gemini model."""
    from unittest.mock import MagicMock, AsyncMock

    model = MagicMock()
    response = MagicMock()
    response.text = "This is a mock AI response."
    model.generate_content = MagicMock(return_value=response)
    model.generate_content_async = AsyncMock(return_value=response)

    return model


# ============================================================================
# Redis Fixtures
# ============================================================================

@pytest.fixture
def mock_redis_client(monkeypatch):
    """Mock Redis client for testing rate limiting."""
    from unittest.mock import MagicMock

    mock_redis = MagicMock()
    mock_redis.get.return_value = None
    mock_redis.setex.return_value = True
    mock_redis.incr.return_value = 1
    mock_redis.ttl.return_value = 60
    mock_redis.delete.return_value = 1

    import middleware.rate_limiter as rate_limiter_module
    monkeypatch.setattr(rate_limiter_module, 'redis_client', mock_redis)

    return mock_redis


# ============================================================================
# Environment Fixtures
# ============================================================================

@pytest.fixture
def test_env_vars(monkeypatch):
    """Set up test environment variables."""
    monkeypatch.setenv('TELEGRAM_BOT_TOKEN', 'test_bot_token_123456')
    monkeypatch.setenv('GEMINI_API_KEY', 'test_gemini_key_123456')
    monkeypatch.setenv('OPENAI_API_KEY', 'test_openai_key_123456')
    monkeypatch.setenv('DB_HOST', 'localhost')
    monkeypatch.setenv('DB_PORT', '5432')
    monkeypatch.setenv('DB_USER', 'test_user')
    monkeypatch.setenv('DB_PASS', 'test_password')
    monkeypatch.setenv('DB_NAME', 'test_db')
    monkeypatch.setenv('REDIS_URL', 'redis://localhost:6379/1')


# ============================================================================
# Cleanup Fixtures
# ============================================================================

@pytest.fixture(autouse=True)
def cleanup_downloads(temp_directory):
    """Cleanup downloaded files after each test."""
    yield
    # Cleanup is handled by temp_directory fixture
