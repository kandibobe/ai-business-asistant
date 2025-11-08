import os
import sys
import fitz  # PyMuPDF
from pydub import AudioSegment
from telegram import Bot, InlineKeyboardButton, InlineKeyboardMarkup
from sqlalchemy.orm import Session
from dotenv import load_dotenv
import pandas as pd
from docx import Document
from openai import OpenAI
import requests
from bs4 import BeautifulSoup
from urllib.parse import urlparse

load_dotenv()

from celery_app import app
from database.database import SessionLocal
from database import crud

# Явно экспортируем все задачи для Celery
__all__ = [
    'process_pdf_task',
    'transcribe_audio_task',
    'process_excel_task',
    'process_word_task',
    'scrape_url_task',
]

TELEGRAM_BOT_TOKEN = os.getenv('TELEGRAM_BOT_TOKEN')
if not TELEGRAM_BOT_TOKEN:
    raise ValueError("TELEGRAM_BOT_TOKEN не найден в .env файле!")
bot = Bot(token=TELEGRAM_BOT_TOKEN)

def get_post_analysis_keyboard() -> InlineKeyboardMarkup:
    """Клавиатура, отправляемая после успешного анализа."""
    keyboard = [
        [InlineKeyboardButton("❓ Задать вопрос по этому документу", switch_inline_query_current_chat="")],
        [InlineKeyboardButton("📚 Показать все мои документы", callback_data='my_docs')],
    ]
    return InlineKeyboardMarkup(keyboard)

@app.task
def process_pdf_task(chat_id: int, user_id: int, username: str, first_name: str, last_name: str, file_path: str, file_name: str):
    """Celery-задача для асинхронной обработки PDF."""
    print(f"WORKER: Начал обработку PDF {file_name}")
    text = ""
    try:
        with fitz.open(file_path) as doc:
            for page in doc:
                text += page.get_text()
    except Exception as e:
        bot.send_message(chat_id, f"❌ Не удалось обработать PDF '{file_name}'. Ошибка: {e}")
        return

    db: Session = SessionLocal()
    try:
        db_user = crud.get_or_create_user(db, user_id, username, first_name, last_name)
        new_doc = crud.create_user_document(db, db_user, file_name, file_path, text)
        # АВТОМАТИЧЕСКИ делаем новый документ активным
        crud.set_active_document(db, db_user, new_doc.id)
        
        bot.send_message(
            chat_id,
            f"✅ PDF '{file_name}' успешно проанализирован и сохранен.\n"
            f"📄 **Он назначен активным для диалога.**\n\n"
            f"Извлечено {len(text)} символов. Что делаем дальше?",
            parse_mode='HTML',
            reply_markup=get_post_analysis_keyboard()
        )
    finally:
        db.close()
    print(f"WORKER: Закончил обработку PDF {file_name}")

@app.task
def transcribe_audio_task(chat_id: int, user_id: int, username: str, first_name: str, last_name: str, file_path: str, file_name: str):
    """Celery-задача для транскрибации аудио с помощью OpenAI Whisper API."""
    print(f"WORKER: Начал транскрибацию {file_name}")

    openai_api_key = os.getenv('OPENAI_API_KEY')

    try:
        if not openai_api_key:
            # Если ключ не настроен, используем заглушку
            print("⚠️ OPENAI_API_KEY не найден, использую заглушку")
            sound = AudioSegment.from_file(file_path)
            text = f"[DEMO MODE] Транскрибация аудио '{file_name}'. Длительность: {len(sound) / 1000:.2f} сек.\n\n"
            text += "Для реальной транскрибации настройте OPENAI_API_KEY в .env файле."
        else:
            # Реальная транскрибация через Whisper API
            client = OpenAI(api_key=openai_api_key)

            with open(file_path, "rb") as audio_file:
                transcript = client.audio.transcriptions.create(
                    model="whisper-1",
                    file=audio_file,
                    language="ru"  # Можно сделать автоопределение, убрав этот параметр
                )
            text = transcript.text

    except Exception as e:
        bot.send_message(chat_id, f"❌ Не удалось обработать аудио '{file_name}'. Ошибка: {e}")
        return

    db: Session = SessionLocal()
    try:
        db_user = crud.get_or_create_user(db, user_id, username, first_name, last_name)
        new_doc = crud.create_user_document(db, db_user, file_name, file_path, text)
        # АВТОМАТИЧЕСКИ делаем новый документ активным
        crud.set_active_document(db, db_user, new_doc.id)

        bot.send_message(
            chat_id,
            f"✅ Аудио '{file_name}' успешно транскрибировано и сохранено.\n"
            f"📄 **Запись назначена активной для диалога.**\n\n"
            f"Распознано {len(text)} символов. Что делаем дальше?",
            parse_mode='HTML',
            reply_markup=get_post_analysis_keyboard()
        )
    finally:
        db.close()
    print(f"WORKER: Закончил транскрибацию {file_name}")

@app.task
def process_excel_task(chat_id: int, user_id: int, username: str, first_name: str, last_name: str, file_path: str, file_name: str):
    """Celery-задача для асинхронной обработки Excel файлов."""
    print(f"WORKER: Начал обработку Excel {file_name}")
    text = ""

    try:
        # Читаем все листы Excel файла
        excel_file = pd.ExcelFile(file_path)

        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name)

            # Добавляем название листа
            text += f"\n{'='*50}\nЛИСТ: {sheet_name}\n{'='*50}\n\n"

            # Конвертируем DataFrame в текст с сохранением структуры
            text += df.to_string(index=False, na_rep='')
            text += "\n\n"

            # Добавляем базовую статистику для числовых столбцов
            numeric_cols = df.select_dtypes(include=['number']).columns
            if len(numeric_cols) > 0:
                text += f"--- Статистика по числовым столбцам ---\n"
                text += df[numeric_cols].describe().to_string()
                text += "\n\n"

        # Добавляем метаинформацию
        text += f"\n{'='*50}\nМЕТАИНФОРМАЦИЯ\n{'='*50}\n"
        text += f"Всего листов: {len(excel_file.sheet_names)}\n"
        text += f"Названия листов: {', '.join(excel_file.sheet_names)}\n"

    except Exception as e:
        bot.send_message(chat_id, f"❌ Не удалось обработать Excel '{file_name}'. Ошибка: {e}")
        return

    db: Session = SessionLocal()
    try:
        db_user = crud.get_or_create_user(db, user_id, username, first_name, last_name)
        new_doc = crud.create_user_document(db, db_user, file_name, file_path, text)
        # АВТОМАТИЧЕСКИ делаем новый документ активным
        crud.set_active_document(db, db_user, new_doc.id)

        bot.send_message(
            chat_id,
            f"✅ Excel файл '{file_name}' успешно проанализирован и сохранен.\n"
            f"📄 **Он назначен активным для диалога.**\n\n"
            f"📊 Обработано листов: {len(excel_file.sheet_names)}\n"
            f"Извлечено {len(text)} символов. Что делаем дальше?",
            parse_mode='HTML',
            reply_markup=get_post_analysis_keyboard()
        )
    finally:
        db.close()
    print(f"WORKER: Закончил обработку Excel {file_name}")

@app.task
def process_word_task(chat_id: int, user_id: int, username: str, first_name: str, last_name: str, file_path: str, file_name: str):
    """Celery-задача для асинхронной обработки Word файлов."""
    print(f"WORKER: Начал обработку Word {file_name}")
    text = ""

    try:
        doc = Document(file_path)

        # Извлекаем текст из параграфов
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"

        # Извлекаем текст из таблиц
        if doc.tables:
            text += "\n" + "="*50 + "\nТАБЛИЦЫ\n" + "="*50 + "\n\n"

            for i, table in enumerate(doc.tables, 1):
                text += f"--- Таблица {i} ---\n"
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    text += row_text + "\n"
                text += "\n"

        # Метаинформация
        text += f"\n{'='*50}\nМЕТАИНФОРМАЦИЯ\n{'='*50}\n"
        text += f"Всего параграфов: {len(doc.paragraphs)}\n"
        text += f"Всего таблиц: {len(doc.tables)}\n"

    except Exception as e:
        bot.send_message(chat_id, f"❌ Не удалось обработать Word '{file_name}'. Ошибка: {e}")
        return

    db: Session = SessionLocal()
    try:
        db_user = crud.get_or_create_user(db, user_id, username, first_name, last_name)
        new_doc = crud.create_user_document(db, db_user, file_name, file_path, text)
        # АВТОМАТИЧЕСКИ делаем новый документ активным
        crud.set_active_document(db, db_user, new_doc.id)

        bot.send_message(
            chat_id,
            f"✅ Word файл '{file_name}' успешно проанализирован и сохранен.\n"
            f"📄 **Он назначен активным для диалога.**\n\n"
            f"📝 Параграфов: {len(doc.paragraphs)} | Таблиц: {len(doc.tables)}\n"
            f"Извлечено {len(text)} символов. Что делаем дальше?",
            parse_mode='HTML',
            reply_markup=get_post_analysis_keyboard()
        )
    finally:
        db.close()
    print(f"WORKER: Закончил обработку Word {file_name}")

@app.task
def scrape_url_task(chat_id: int, user_id: int, username: str, first_name: str, last_name: str, url: str):
    """Celery-задача для веб-скрапинга и анализа URL."""
    print(f"WORKER: Начал скрапинг URL {url}")
    text = ""

    try:
        # Проверяем валидность URL
        parsed_url = urlparse(url)
        if not all([parsed_url.scheme, parsed_url.netloc]):
            bot.send_message(chat_id, f"❌ Некорректный URL: {url}")
            return

        # Делаем запрос к URL
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        }
        response = requests.get(url, headers=headers, timeout=30)
        response.raise_for_status()

        # Парсим HTML
        soup = BeautifulSoup(response.content, 'html.parser')

        # Извлекаем title
        title = soup.find('title')
        if title:
            text += f"ЗАГОЛОВОК СТРАНИЦЫ:\n{title.get_text()}\n\n"

        # Извлекаем meta description
        meta_desc = soup.find('meta', attrs={'name': 'description'})
        if meta_desc and meta_desc.get('content'):
            text += f"ОПИСАНИЕ:\n{meta_desc.get('content')}\n\n"

        # Извлекаем основной текст
        text += "="*50 + "\nОСНОВНОЙ КОНТЕНТ\n" + "="*50 + "\n\n"

        # Удаляем скрипты и стили
        for script in soup(['script', 'style', 'nav', 'footer', 'header']):
            script.decompose()

        # Извлекаем текст из основных тегов
        main_content = soup.find('main') or soup.find('article') or soup.find('body')
        if main_content:
            paragraphs = main_content.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'li'])
            for p in paragraphs:
                p_text = p.get_text().strip()
                if p_text:
                    text += p_text + "\n"

        # Метаинформация
        text += f"\n{'='*50}\nМЕТАИНФОРМАЦИЯ\n{'='*50}\n"
        text += f"URL: {url}\n"
        text += f"Домен: {parsed_url.netloc}\n"
        text += f"Длина контента: {len(response.content)} байт\n"

    except requests.exceptions.RequestException as e:
        bot.send_message(chat_id, f"❌ Не удалось загрузить URL '{url}'. Ошибка: {e}")
        return
    except Exception as e:
        bot.send_message(chat_id, f"❌ Ошибка при обработке URL '{url}'. Ошибка: {e}")
        return

    db: Session = SessionLocal()
    try:
        db_user = crud.get_or_create_user(db, user_id, username, first_name, last_name)
        # Сохраняем как "документ" с именем URL
        file_name = f"Веб-страница: {parsed_url.netloc}"
        new_doc = crud.create_user_document(db, db_user, file_name, url, text)
        # АВТОМАТИЧЕСКИ делаем новый документ активным
        crud.set_active_document(db, db_user, new_doc.id)

        bot.send_message(
            chat_id,
            f"✅ Веб-страница '{parsed_url.netloc}' успешно проанализирована.\n"
            f"📄 **Страница назначена активной для диалога.**\n\n"
            f"🌐 URL: {url}\n"
            f"Извлечено {len(text)} символов. Что делаем дальше?",
            parse_mode='HTML',
            reply_markup=get_post_analysis_keyboard()
        )
    finally:
        db.close()
    print(f"WORKER: Закончил скрапинг URL {url}")